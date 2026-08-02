"""合同生成引擎 - 模板 + LLM生成 → DOCX导出"""
import json
import time
import urllib.request
import urllib.error
import os
from datetime import datetime
from config import LLM_API_KEY, LLM_BASE_URL, LLM_MODEL

_CHECKLIST_PATH = os.path.join(os.path.dirname(os.path.dirname(__file__)), "data", "avoidance_checklist.txt")


def _load_checklist() -> str:
    try:
        if os.path.exists(_CHECKLIST_PATH):
            return open(_CHECKLIST_PATH, encoding="utf-8").read().strip()
    except OSError:
        return ""
    return ""


def _append_checklist(text: str):
    if not text.strip():
        return
    try:
        os.makedirs(os.path.dirname(_CHECKLIST_PATH), exist_ok=True)
        existing = set()
        if os.path.exists(_CHECKLIST_PATH):
            existing = {line.strip() for line in open(_CHECKLIST_PATH, encoding="utf-8").readlines() if line.strip()}
        new_lines = [l.strip() for l in text.split("\n") if l.strip() and l.strip() not in existing]
        if new_lines:
            with open(_CHECKLIST_PATH, "a", encoding="utf-8") as f:
                f.write("\n".join(new_lines) + "\n")
    except OSError:
        pass


# ponytail: public alias so app.py doesn't import a private fn
append_checklist = _append_checklist

_CONTRACT_PROMPTS = {
    "采购合同": """你是合同起草专家。请根据以下要求起草一份采购合同。

要求：{requirements}
买方：{buyer}；卖方：{seller}

请写出完整合同，包含以下条款：
1. 买卖双方信息；2. 采购标的（品名、规格、数量）；3. 价款与支付方式；
4. 交付时间、地点和方式；5. 验收标准与程序；6. 违约责任；
7. 争议解决；8. 合同生效与变更；9. 不可抗力；10. 其他约定。

输出完整合同文本，法定条款缺一不可。""",

    "服务合同": """你是合同起草专家。请根据以下要求起草一份服务合同。

要求：{requirements}
委托方：{buyer}；服务方：{seller}

请写出完整合同，包含以下条款：
1. 双方信息；2. 服务内容与范围；3. 服务期限；4. 服务费用与支付；
5. 双方权利义务；6. 验收标准；7. 保密义务；8. 知识产权归属；
9. 违约责任；10. 争议解决；11. 合同变更与终止。

输出完整合同文本。""",

    "劳动合同": """你是合同起草专家。请根据以下要求起草一份劳动合同。

要求：{requirements}
用人单位：{buyer}；劳动者：{seller}

请写出完整合同，包含以下条款：
1. 合同期限（含试用期）；2. 工作内容与地点；3. 工作时间与休息休假；
4. 劳动报酬与社会保险；5. 劳动保护；6. 合同解除与终止；7. 竞业限制与保密；
8. 违约责任；9. 争议解决。

严格遵循《劳动合同法》要求。输出完整合同文本。""",

    "租赁合同": """你是合同起草专家。请根据以下要求起草一份租赁合同。

要求：{requirements}
出租方：{buyer}；承租方：{seller}

请写出完整合同，包含以下条款：
1. 租赁物描述；2. 租赁期限；3. 租金与支付方式；4. 押金；
5. 使用与维护；6. 转租限制；7. 合同解除；8. 违约责任；9. 争议解决。

输出完整合同文本。""",

    "保密协议(NDA)": """你是合同起草专家。请根据以下要求起草一份保密协议。

要求：{requirements}
披露方：{buyer}；接收方：{seller}

请写出完整协议，包含以下条款：
1. 保密信息定义；2. 保密义务范围；3. 保密期限；
4. 例外情况；5. 违约责任；6. 争议解决。

输出完整协议文本。""",

    "合作框架协议": """你是合同起草专家。请根据以下要求起草一份合作框架协议。

要求：{requirements}
甲方：{buyer}；乙方：{seller}

请写出完整协议，包含以下条款：
1. 合作背景与目的；2. 合作内容与方式；3. 双方权利义务；
4. 资源投入与利益分配；5. 知识产权；6. 保密义务；
7. 合作期限与终止；8. 违约责任；9. 争议解决。

输出完整协议文本。""",

    "销售合同": """你是合同起草专家。请根据以下要求起草一份销售合同。

要求：{requirements}
买方：{buyer}；卖方：{seller}

请写出完整合同，包含以下条款：
1. 产品信息（名称、规格、数量、单价）；2. 质量要求与标准；
3. 包装要求；4. 交付与运输；5. 验收；6. 付款方式与期限；
7. 售后服务；8. 违约责任；9. 争议解决。

输出完整合同文本。""",

    "技术开发合同": """你是合同起草专家。请根据以下要求起草一份技术开发合同。

要求：{requirements}
委托方：{buyer}；开发方：{seller}

请写出完整合同，包含以下条款：
1. 开发内容与技术指标；2. 开发周期与里程碑；3. 开发费用与支付节点；
4. 验收标准与程序；5. 知识产权归属与使用许可；6. 技术支持与维护；
7. 保密义务；8. 违约责任；9. 争议解决。

输出完整合同文本。""",
}


def _sanitize_format_input(text):
    return text.replace('{', '{{').replace('}', '}}')


def generate_contract(contract_type: str, requirements: str,
                      buyer: str = "", seller: str = "",
                      review_feedback: str = "",
                      load_checklist: bool = True) -> dict:
    if contract_type not in _CONTRACT_PROMPTS:
        return {"error": f"不支持的合同类型: {contract_type}", "supported": list(_CONTRACT_PROMPTS.keys())}

    template = _CONTRACT_PROMPTS[contract_type]
    safe_requirements = _sanitize_format_input(requirements)
    safe_buyer = _sanitize_format_input(buyer or "甲方")
    safe_seller = _sanitize_format_input(seller or "乙方")
    prompt = template.format(
        requirements=safe_requirements,
        buyer=safe_buyer,
        seller=safe_seller
    )
    prompt = prompt.replace('{{', '{').replace('}}', '}')

    # ponytail: inject accumulated avoidance checklist + review feedback
    if load_checklist:
        checklist = _load_checklist()
        if checklist:
            prompt += "\n\n【历史避坑规则 — 请严格遵守，避免以下已知问题】\n" + checklist
    if review_feedback:
        prompt += "\n" + review_feedback
    # 强制把需求中的具体数字直接填入条款，禁止留空白待填
    prompt += "\n\n【硬性要求】需求描述中出现的具体数字（金额、期限、数量、比例、天数等）必须直接写入对应条款，不得留空白或提示用户自行填写。"

    if not LLM_API_KEY:
        fallback = _fallback_template(contract_type, buyer, seller)
        return {"generated_text": fallback,
                "contract_type": contract_type, "word_count": len(fallback),
                "disclaimer": "本合同时由模板生成(LLM未配置)，仅供参考，正式签署前请由专业律师审核。",
                "llm_available": False}

    text = _call_llm(prompt)
    if text is None:
        fallback = _fallback_template(contract_type, buyer, seller)
        return {"generated_text": fallback,
                "contract_type": contract_type, "word_count": len(fallback),
                "disclaimer": "LLM调用失败，已降级为模板生成。正式签署前请由专业律师审核。",
                "llm_available": False}

    return {
        "contract_type": contract_type,
        "generated_text": text,
        "word_count": len(text),
        "disclaimer": "本合同时由AI生成，仅供参考，正式签署前请由专业律师审核。",
        "llm_available": True
    }


def export_docx(text: str, contract_type: str, output_path: str = "") -> str:
    try:
        from docx import Document
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return ""

    doc = Document()

    title = doc.add_heading(f'{contract_type}', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for para_text in text.split('\n\n'):
        para_text = para_text.strip()
        if not para_text:
            continue
        if len(para_text) < 40 and ('第' in para_text or '条' in para_text or
                                      para_text.startswith(('一', '二', '三', '四', '五', '六', '七', '八', '九', '十'))):
            doc.add_heading(para_text, level=2)
        else:
            p = doc.add_paragraph()
            run = p.add_run(para_text)
            run.font.size = Pt(11)

    doc.add_paragraph('')
    footer = doc.add_paragraph()
    footer_run = footer.add_run('【免责声明】本合同时由AI智能生成，仅供参考。正式签署前请由执业律师审核。')
    footer_run.font.size = Pt(9)

    if not output_path:
        ts = datetime.now().strftime('%Y%m%d_%H%M%S')
        output_path = os.path.join(os.path.dirname(os.path.dirname(__file__)),
                                   'uploads', f'{contract_type}_{ts}.docx')

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return output_path


def _call_llm(prompt: str, max_retries: int = 2) -> str:
    data = json.dumps({
        "model": LLM_MODEL,
        "messages": [{"role": "user", "content": prompt}],
        "temperature": 0.7, "max_tokens": 4000
    }).encode()

    req = urllib.request.Request(
        f"{LLM_BASE_URL.rstrip('/')}/chat/completions",
        data=data,
        headers={"Authorization": f"Bearer {LLM_API_KEY}", "Content-Type": "application/json"}
    )

    for attempt in range(max_retries + 1):
        try:
            with urllib.request.urlopen(req, timeout=60) as resp:
                result = json.loads(resp.read())
                return result["choices"][0]["message"]["content"]
        except urllib.error.HTTPError as e:
            if e.code == 429 and attempt < max_retries:
                time.sleep(2 ** attempt)
                continue
            return None
        except (urllib.error.URLError, json.JSONDecodeError, KeyError, TimeoutError) as e:
            if attempt < max_retries:
                time.sleep(2 ** attempt)
                continue
            return None
    return None


def _fallback_template(contract_type: str, buyer: str, seller: str) -> str:
    buyer = buyer or "甲方"
    seller = seller or "乙方"
    return f"""# {contract_type}

## 一、当事人信息

甲方（{buyer}）：
地址：
法定代表人：
联系电话：

乙方（{seller}）：
地址：
法定代表人：
联系电话：

## 二、合同主要条款

（本合同时由模板生成，建议配置LLM API Key以获得完整合同文本）

## 三、违约责任

任何一方违反本合同约定，应承担相应的违约责任。

## 四、争议解决

因本合同引起的争议，双方协商解决；协商不成的，向合同签订地人民法院提起诉讼。

## 五、其他

本合同一式两份，双方各执一份，具有同等法律效力。
签署日期：{datetime.now().strftime('%Y年%m月%d日')}
"""
