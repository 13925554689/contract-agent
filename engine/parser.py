"""合同解析引擎 - PDF/DOCX/TXT → 结构化条款"""
import re

from pathlib import Path
try:
    from PyPDF2 import PdfReader
except ImportError:
    PdfReader = None
try:
    from docx import Document
except ImportError:
    Document = None


def parse_contract(filepath: str) -> dict:
    """解析合同文件，返回结构化文本和元信息"""
    ext = Path(filepath).suffix.lower()

    if ext == '.pdf':
        text = _parse_pdf(filepath)
    elif ext in ('.docx', '.doc'):
        text = _parse_docx(filepath)
    elif ext == '.txt':
        text = Path(filepath).read_text(encoding='utf-8', errors='ignore')
    else:
        raise ValueError(f"不支持的文件格式: {ext}")

    # 清理文本
    text = _clean_text(text)

    # 结构化拆分
    metadata = _extract_metadata(text)
    clauses = _split_clauses(text)

    return {
        "raw_text": text,
        "metadata": metadata,
        "clauses": clauses,
        "clause_count": len(clauses),
        "word_count": len(text)
    }


def _clean_text(text: str) -> str:
    """清理文本噪声"""
    text = re.sub(r' +', ' ', text)           # 合并空格
    text = re.sub(r'\n{3,}', '\n\n', text)    # 合并空行
    text = re.sub(r'(\w)-\n(\w)', r'\1\2', text)  # 修复断词
    return text.strip()


def _parse_pdf(filepath: str) -> str:
    if PdfReader is None:
        raise ImportError("PyPDF2未安装")
    reader = PdfReader(filepath)
    text_parts = []
    for page in reader.pages:
        t = page.extract_text()
        if t:
            text_parts.append(t)
    return '\n'.join(text_parts)


def _parse_docx(filepath: str) -> str:
    if Document is None:
        raise ImportError("python-docx未安装")
    doc = Document(filepath)
    text_parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)

    # 表格提取
    for table in doc.tables:
        text_parts.append('---[表格]---')
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                text_parts.append(' | '.join(cells))

    return '\n'.join(text_parts)


def _extract_metadata(text: str) -> dict:
    """提取合同元信息"""
    info = {
        "title": "", "party_a": "", "party_b": "",
        "contract_amount": "", "effective_date": "", "expiry_date": "",
        "governing_law": "", "contract_type": ""
    }

    # 合同标题
    m = re.search(r'《?(.{2,30}(?:合[同約约]|协[议議]|书[面]?|契[约約]))》?', text[:500])
    if m:
        info["title"] = m.group(1)

    # 甲乙方
    m = re.search(r'(?:甲[方方]|卖[方方]|出[让讓][方方]|发[包包][方方])[：:：\s]*([^\n]{2,30})', text[:1000])
    if m:
        info["party_a"] = m.group(1).strip()
    m = re.search(r'(?:乙[方方]|买[方方]|受[让讓][方方]|承[包包][方方])(?:[（(].*?[）)])?[：:：\s]*([^\n]{2,30})', text[:1000])
    if m:
        info["party_b"] = m.group(1).strip()

    # 金额
    m = re.search(r'(?:总[价價]?[金金额]|合[同約约][总總]?[价價]?[金金额]|人民币).*?([¥￥]?\s*\d[\d,.]+\s*(?:万|元|美元))', text[:2000])
    if m:
        info["contract_amount"] = m.group(1)

    # 日期
    dates = re.findall(r'(\d{4}[年/-]\d{1,2}[月/-]\d{1,2}[日号]?)', text[:2000])
    if dates:
        info["effective_date"] = dates[0]
        if len(dates) > 1:
            info["expiry_date"] = dates[-1]

    # 管辖法
    m = re.search(r'(?:管辖|适用|争议).*?(法[律院])', text[:3000])
    if m:
        info["governing_law"] = m.group(0)[:40]

    return info


def _split_clauses(text: str) -> list:
    """将文本拆分为条款列表"""
    clauses = []

    # 按常见条款标识拆分
    pattern = r'(?:(?:第[一二三四五六七八九十百千0-9]+[条節节章]|(?:一|二|三|四|五|六|七|八|九|十)[、，,.])\s*|(?:^|\n)\s*\d+[\.、．]\s+)'
    parts = re.split(pattern, text)

    if len(parts) <= 2:
        # 简单按段落拆分
        paras = [p.strip() for p in text.split('\n\n') if len(p.strip()) > 10]
        for i, para in enumerate(paras):
            clause_type = _classify_clause(para)
            clauses.append({
                "index": i + 1, "type": clause_type,
                "content": para[:1000], "word_count": len(para)
            })
    else:
        clause_idx = 0
        for i, part in enumerate(parts):
            part = part.strip()
            if len(part) < 10:
                continue
            clause_idx += 1
            clause_type = _classify_clause(part)
            clauses.append({
                "index": clause_idx, "type": clause_type,
                "content": part[:1000], "word_count": len(part)
            })

    return clauses


def _classify_clause(text: str) -> str:
    """自动分类条款类型"""
    # 优先从标题行判断
    first_line = text.split('\n')[0].strip()
    # ponytail: title patterns have higher priority than body keyword matching
    title_patterns = [
        (r'(?:违约|罚[则款]|赔偿)', "违约责任"),
        (r'(?:付[款金]|支付|价款|报酬|费用结算|结算)', "付款条款"),
        (r'(?:交付|交货|验收|质量标准)', "交付条款"),
        (r'(?:保密|机密|商业秘密|NDA)', "保密条款"),
        (r'(?:期限|有效期|合同期限|起止)', "期限条款"),
        (r'(?:知识产权|专利|著作权|商标)', "知识产权"),
        (r'(?:争议|仲裁|诉讼|管辖)', "争议解决"),
        (r'(?:解除|终止|不可抗力)', "解除终止"),
        (r'(?:SLA|服务级别|可用性|响应时间)', "服务等级"),
        (r'(?:数据|个人信息|隐私|安全)', "数据保护"),
    ]
    for pattern, ctype in title_patterns:
        if re.search(pattern, first_line):
            return ctype

    # 回退全文关键词
    body_patterns = {
        "付款条款": [r'付款', r'支付', r'价款', r'报酬', r'费用结算'],
        "违约责任": [r'违约', r'罚[则款]', r'赔偿', r'赔[付偿]'],
        "交付条款": [r'交付', r'交货', r'验收', r'质量标准', r'验收标准'],
        "保密条款": [r'保密', r'机密', r'商业秘密', r'NDA'],
        "期限条款": [r'期限', r'有效期', r'合同期限', r'起止'],
        "知识产权": [r'知识产权', r'专利', r'著作权', r'商标', r'IP'],
        "争议解决": [r'争议', r'仲裁', r'诉讼', r'管辖'],
        "解除终止": [r'解除', r'终止', r'不可抗力'],
        "服务等级": [r'SLA', r'服务级别', r'可用性', r'响应时间'],
        "数据保护": [r'数据', r'个人信息', r'隐私', r'GDPR', r'安全'],
    }
    for clause_type, keywords in body_patterns.items():
        for kw in keywords:
            if kw in text:
                return clause_type

    return "一般条款"
